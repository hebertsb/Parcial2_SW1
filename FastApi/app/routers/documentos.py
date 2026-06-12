"""
NexusFlow AI — Router de Gestión Documental (Segundo Parcial)
=============================================================
Endpoints para subir, listar y descargar documentos de trámites.
Usa MinIO (local) o AWS S3 (producción) — mismo código, diferente endpoint.

Flujo de un documento en un trámite:
  Cliente sube PDF → S3/MinIO guarda en empresa/política/trámite/
  Empleado lista los docs del trámite → ve qué falta
  Empleado descarga un doc específico → URL pre-firmada temporal
"""

import logging
from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Form, status
from fastapi.responses import JSONResponse, StreamingResponse

from ..services.s3_service import get_s3_service
from ..config import settings

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/documentos", tags=["Gestión Documental - S3"])


def _resolver_nombres_doc(politica_id: str, tramite_id: str) -> tuple[str, str]:
    """Resuelve politica_id → nombre_política y tramite_id → nombre_cliente para la ruta en MinIO."""
    try:
        from pymongo import MongoClient
        from bson import ObjectId

        client = MongoClient(settings.mongodb_uri, serverSelectionTimeoutMS=3000)
        db = client[settings.mongodb_database]

        # ── Nombre de la política ─────────────────────────────────
        pol_nombre = ""
        for col in ("Politica", "politicas"):
            p = db[col].find_one({"_id": politica_id}, {"nombre": 1})
            if not p:
                try:
                    p = db[col].find_one({"_id": ObjectId(politica_id)}, {"nombre": 1})
                except Exception:
                    pass
            if p and p.get("nombre"):
                pol_nombre = p["nombre"]
                break

        # ── Nombre del cliente (via tramite → cliente_id → usuario) ──
        cli_nombre = ""
        for col in ("Tramite", "tramites"):
            t = db[col].find_one({"_id": tramite_id}, {"cliente_id": 1})
            if not t:
                try:
                    t = db[col].find_one({"_id": ObjectId(tramite_id)}, {"cliente_id": 1})
                except Exception:
                    pass
            if t and t.get("cliente_id"):
                try:
                    u = db.Usuario.find_one(
                        {"_id": ObjectId(str(t["cliente_id"]))},
                        {"nombre_completo": 1},
                    )
                    if u and u.get("nombre_completo"):
                        cli_nombre = u["nombre_completo"]
                except Exception:
                    pass
                break

        client.close()
        return pol_nombre, cli_nombre
    except Exception as exc:
        logger.warning("_resolver_nombres_doc error: %s", exc)
        return "", ""


@router.post(
    "/subir",
    summary="Subir documento al trámite",
    description=(
        "Sube un archivo (PDF, Word, Excel, imagen) al repositorio S3 del trámite. "
        "Cada trámite tiene su propio espacio aislado por empresa/política/trámite. "
        "Retorna la key S3 y URL de descarga temporal (válida 1 hora)."
    )
)
async def subir_documento(
    archivo: UploadFile = File(..., description="Archivo a subir"),
    empresa_id: str = Form(..., description="ID de la empresa"),
    politica_id: str = Form(..., description="ID de la política de negocio"),
    tramite_id: str = Form(..., description="ID del trámite"),
    subido_por: str = Form("cliente", description="ID o nombre del usuario que sube"),
):
    """
    Sube un documento al repositorio del trámite en S3/MinIO.

    **Formatos soportados:** PDF, Word (.docx), Excel (.xlsx), imágenes (JPG/PNG), video (MP4)

    **Estructura en S3:**
    ```
    nexusflow-documentos/
    └── {empresa_id}/{politica_id}/{tramite_id}/{nombre_archivo}
    ```
    """
    if not archivo.filename:
        raise HTTPException(status_code=400, detail="Archivo sin nombre")

    # Validar tamaño (máximo 50 MB)
    contenido = await archivo.read()
    if len(contenido) > 50 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Archivo demasiado grande (máximo 50 MB)")

    pol_nombre, cli_nombre = _resolver_nombres_doc(politica_id, tramite_id)

    svc = get_s3_service()
    resultado = svc.subir_archivo(
        contenido=contenido,
        nombre_archivo=archivo.filename,
        empresa_id=empresa_id,
        politica_id=politica_id,
        tramite_id=tramite_id,
        subido_por=subido_por,
        politica_nombre=pol_nombre,
        cliente_nombre=cli_nombre,
    )

    if "error" in resultado:
        raise HTTPException(status_code=500, detail=resultado["error"])

    return resultado


@router.get(
    "/excel-data/{key:path}",
    summary="Convertir Excel a JSON para el visor de planilla",
    description="Lee un archivo .xlsx desde MinIO y retorna todos los datos de hojas en formato JSON para el visor Angular."
)
async def excel_datos_json(key: str):
    """Convierte XLSX a JSON estructurado: hojas, filas, celdas con estilos y anchos de columna."""
    svc = get_s3_service()
    resultado = svc.descargar_archivo(key)
    if "error" in resultado:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    try:
        import openpyxl
        from openpyxl.utils import get_column_letter
        from io import BytesIO

        contenido_bytes = resultado["body"].read()
        wb = openpyxl.load_workbook(BytesIO(contenido_bytes), data_only=True)

        THEME_COLORS = {
            0: 'FFFFFF', 1: '000000', 2: 'E7E6E6', 3: '44546A',
            4: '4472C4', 5: 'ED7D31', 6: 'A9D18E', 7: 'FFC000',
            8: '5B9BD5', 9: '70AD47',
        }

        def get_color_hex(color_obj):
            if color_obj is None:
                return None
            try:
                if color_obj.type == 'rgb':
                    rgb = color_obj.rgb or ''
                    if len(rgb) == 8:
                        if rgb[:2] == '00':
                            return None
                        hex_color = rgb[2:]
                        if hex_color.upper() in ('FFFFFF', '000000'):
                            return None
                        return f"#{hex_color}"
                    elif len(rgb) == 6 and rgb.upper() not in ('FFFFFF', '000000'):
                        return f"#{rgb}"
                elif color_obj.type == 'theme':
                    base_hex = THEME_COLORS.get(color_obj.theme or 0, 'FFFFFF')
                    if base_hex in ('FFFFFF', '000000'):
                        return None
                    return f"#{base_hex}"
            except Exception:
                pass
            return None

        sheets = []
        for ws in wb.worksheets:
            col_widths_raw: dict = {}
            for col_letter, col_dim in ws.column_dimensions.items():
                if col_dim.width:
                    col_widths_raw[col_letter] = max(50, int(col_dim.width * 7))

            max_row = ws.max_row or 0
            max_col = ws.max_column or 0

            if max_row == 0 or max_col == 0:
                sheets.append({"name": ws.title, "rows": [], "col_widths": [], "max_row": 0, "max_col": 0})
                continue

            rows = []
            for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
                row_data = []
                for cell in row:
                    val = cell.value
                    if val is None:
                        val = ""
                    elif isinstance(val, float):
                        val = f"{val:g}"
                    elif isinstance(val, bool):
                        val = "VERDADERO" if val else "FALSO"
                    else:
                        val = str(val)

                    cell_info: dict = {"v": val}
                    try:
                        if cell.font:
                            if cell.font.bold:
                                cell_info["bold"] = True
                            fg = get_color_hex(cell.font.color) if cell.font.color else None
                            if fg:
                                cell_info["color"] = fg
                        if cell.fill and cell.fill.fgColor:
                            bg = get_color_hex(cell.fill.fgColor)
                            if bg:
                                cell_info["bg"] = bg
                        if cell.alignment:
                            h = cell.alignment.horizontal
                            if h in ('center', 'right'):
                                cell_info["align"] = h
                    except Exception:
                        pass

                    row_data.append(cell_info)
                rows.append(row_data)

            col_widths = [col_widths_raw.get(get_column_letter(i + 1), 80) for i in range(max_col)]

            sheets.append({
                "name": ws.title,
                "rows": rows,
                "col_widths": col_widths,
                "max_row": max_row,
                "max_col": max_col,
            })

        return {"sheets": sheets}

    except Exception as e:
        logger.error("Error leyendo excel-data: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@router.get(
    "/{empresa_id}/{politica_id}/{tramite_id}",
    summary="Listar documentos del trámite",
    description="Retorna la lista de todos los documentos subidos para un trámite específico, con URLs de descarga."
)
async def listar_documentos(empresa_id: str, politica_id: str, tramite_id: str):
    """
    Lista todos los documentos de un trámite.
    Busca en nueva estructura (nombre_política/nombre_cliente) y estructura anterior (IDs).
    """
    pol_nombre, cli_nombre = _resolver_nombres_doc(politica_id, tramite_id)
    svc = get_s3_service()
    archivos = svc.listar_archivos(
        empresa_id, politica_id, tramite_id,
        politica_nombre=pol_nombre,
        cliente_nombre=cli_nombre,
    )
    return {
        "tramite_id": tramite_id,
        "politica_id": politica_id,
        "empresa_id": empresa_id,
        "total_documentos": len(archivos),
        "documentos": archivos,
    }


@router.get(
    "/url/{key:path}",
    summary="Generar URL de descarga",
    description="Genera una URL pre-firmada para descargar un documento directamente desde S3/MinIO. Válida 1 hora."
)
async def url_descarga(key: str, expira: int = 3600):
    """Genera URL temporal para descargar el documento sin autenticación adicional."""
    svc = get_s3_service()
    url = svc.generar_url_descarga(key, expira)
    if not url:
        raise HTTPException(status_code=404, detail="Documento no encontrado o S3 no disponible")
    return {"url": url, "expira_segundos": expira, "key": key}


@router.get(
    "/proxy/{key:path}",
    summary="Proxy para descargar documentos desde MinIO",
    description="Usado por OnlyOffice (vía red Docker interna) y por el browser para descargas directas."
)
async def proxy_documento(key: str):
    """
    Descarga el archivo desde MinIO y lo retorna al cliente.
    OnlyOffice usa http://fastapi-service:8000/documentos/proxy/{key} (red Docker).
    Browser usa http://localhost:8000/documentos/proxy/{key}.
    """
    svc = get_s3_service()
    resultado = svc.descargar_archivo(key)
    if "error" in resultado:
        raise HTTPException(status_code=404, detail=resultado["error"])
    ct = resultado.get("content_type", "")
    nombre = resultado["nombre"]
    inline_types = ("application/pdf", "image/")
    disposition = "inline" if any(ct.startswith(t) for t in inline_types) else "attachment"
    return StreamingResponse(
        resultado["body"],
        media_type=ct,
        headers={"Content-Disposition": f'{disposition}; filename="{nombre}"'}
    )


@router.post(
    "/onlyoffice-callback",
    summary="Callback de OnlyOffice para guardar documentos",
    description="OnlyOffice llama a este endpoint cuando el usuario guarda. Retorna error=0 para confirmar recepción."
)
async def onlyoffice_callback(request: Request):
    """
    Endpoint requerido por OnlyOffice Document Server.
    status=2 → el documento fue cerrado, OnlyOffice envía el archivo editado.
    Para este demo solo confirmamos recepción (error=0).
    """
    try:
        body = await request.json()
        logger.info("OnlyOffice callback status=%s key=%s", body.get("status"), body.get("key"))
    except Exception:
        pass
    return JSONResponse({"error": 0})


@router.delete(
    "/{key:path}",
    summary="Eliminar documento",
    description="Elimina un documento del repositorio S3/MinIO por su key."
)
async def eliminar_documento(key: str):
    svc = get_s3_service()
    resultado = svc.eliminar_archivo(key)
    if "error" in resultado:
        raise HTTPException(status_code=500, detail=resultado["error"])
    return resultado


@router.get(
    "/leer/{key:path}",
    summary="Leer documento y convertir a HTML para editar",
    description="Lee un documento (.docx, .txt, .html) y lo convierte a HTML para editar en TipTap"
)
async def leer_documento_para_editar(key: str):
    """
    Lee archivo desde MinIO y retorna contenido como HTML.
    - .docx → convierte a HTML
    - .txt → envuelve en <p>
    - .html → retorna como está
    """
    svc = get_s3_service()
    resultado = svc.descargar_archivo(key)

    if "error" in resultado:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    try:
        contenido_bytes = resultado["body"].read()
        nombre = resultado["nombre"]
        ext = nombre.split('.')[-1].lower()

        html = ""

        if ext == "docx":
            import mammoth
            from io import BytesIO
            import base64
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from bs4 import BeautifulSoup

            # Mammoth style_map: handles Spanish Word heading names + English defaults
            _STYLE_MAP = """
p[style-name='Heading 1'] => h1:fresh
p[style-name='Heading 2'] => h2:fresh
p[style-name='Heading 3'] => h3:fresh
p[style-name='Heading 4'] => h4:fresh
p[style-name='Heading 5'] => h5:fresh
p[style-name='Título 1'] => h1:fresh
p[style-name='Título 2'] => h2:fresh
p[style-name='Título 3'] => h3:fresh
p[style-name='Título 4'] => h4:fresh
p[style-name='Title'] => h1:fresh
p[style-name='Subtitle'] => h2:fresh
p[style-name='Section Title'] => h2:fresh
r[style-name='Strong'] => strong
"""

            try:
                def _img_handler(image):
                    try:
                        with image.open() as f:
                            img_bytes = f.read()
                        mime = image.content_type or 'image/png'
                        # Skip WMF/EMF — browsers can't render Windows metafiles
                        if any(x in mime.lower() for x in ('wmf', 'emf', 'x-ms')):
                            return {}
                        b64 = base64.b64encode(img_bytes).decode('utf-8')
                        return {'src': f'data:{mime};base64,{b64}',
                                'style': 'max-width:100%;height:auto;display:block;margin:8px auto;'}
                    except Exception:
                        return {}

                resultado_mammoth = mammoth.convert_to_html(
                    BytesIO(contenido_bytes),
                    convert_image=mammoth.images.img_element(_img_handler),
                    style_map=_STYLE_MAP
                )
                html = resultado_mammoth.value or "<p></p>"

                # Inyectar alineación usando python-docx + mapeo por texto
                try:
                    _ALIGN_CSS = {
                        WD_ALIGN_PARAGRAPH.CENTER: 'center',
                        WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                        WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
                    }

                    def _effective_align(para):
                        """Lee alineación directa o heredada del estilo."""
                        a = para.alignment
                        if a is not None:
                            return a
                        try:
                            style = para.style
                            while style:
                                if style.paragraph_format.alignment is not None:
                                    return style.paragraph_format.alignment
                                style = style.base_style
                        except Exception:
                            pass
                        return None

                    doc_meta = Document(BytesIO(contenido_bytes))
                    # Mapear texto (primeros 100 chars) → alineación CSS
                    align_map: dict[str, str] = {}
                    for para in doc_meta.paragraphs:
                        key = para.text.strip()[:100]
                        if key:
                            css = _ALIGN_CSS.get(_effective_align(para))
                            if css:
                                align_map[key] = css

                    soup = BeautifulSoup(html, 'html.parser')
                    for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                        key = tag.get_text(strip=True)[:100]
                        if key and key in align_map:
                            existing = tag.get('style', '')
                            new_s = f"text-align:{align_map[key]}"
                            tag['style'] = f"{new_s};{existing}" if existing else new_s
                    html = str(soup)
                except Exception as align_err:
                    logger.warning("alignment injection skipped: %s", align_err)

                # Intentar leer dimensiones reales de imágenes del docx
                try:
                    WD_DRAW_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
                    img_widths = []
                    for para in doc_meta.paragraphs:
                        for inline in para._element.findall(f'.//{WD_DRAW_NS}inline'):
                            extent = inline.find(f'{WD_DRAW_NS}extent')
                            if extent is not None:
                                cx = int(extent.get('cx', 0))
                                width_px = int(cx / 914400 * 96)
                                img_widths.append(width_px)

                    if img_widths:
                        soup2 = BeautifulSoup(html, 'html.parser')
                        for i, img_tag in enumerate(soup2.find_all('img')):
                            if i < len(img_widths):
                                w = min(img_widths[i], 600)
                                img_tag['style'] = f"width:{w}px;height:auto;display:block;margin:0 auto;"
                            else:
                                img_tag['style'] = "max-width:50%;height:auto;display:block;margin:0 auto;"
                        html = str(soup2)
                except Exception:
                    pass

                # No <style> tag — TipTap/ProseMirror strips it and it can corrupt HTML parsing
                # Image sizing is handled via inline style in _img_handler + frontend CSS
            except Exception as e:
                logger.error("mammoth conversion failed: %s", e)
                html = "<p>No se pudo convertir el documento .docx. Verifica que el archivo no esté corrupto.</p>"

        elif ext in ("xlsx", "xls"):
            import openpyxl
            from openpyxl.styles.colors import Color
            from io import BytesIO

            wb = openpyxl.load_workbook(BytesIO(contenido_bytes), data_only=True)
            ws = wb.active

            # Colores de tema Excel comunes (Office theme)
            THEME_COLORS = {
                0: 'FFFFFF', 1: '000000', 2: 'E7E6E6', 3: '44546A',
                4: '4472C4', 5: 'ED7D31', 6: 'A9D18E', 7: 'FFC000',
                8: '5B9BD5', 9: '70AD47',
            }

            def get_color_hex(color_obj):
                if color_obj is None: return None
                try:
                    if color_obj.type == 'rgb':
                        rgb = color_obj.rgb or ''
                        if len(rgb) == 8:
                            alpha = rgb[:2]
                            if alpha == '00': return None
                            hex_color = rgb[2:]
                            if hex_color.upper() in ('FFFFFF', '000000'): return None
                            return f"#{hex_color}"
                        elif len(rgb) == 6 and rgb.upper() not in ('FFFFFF', '000000'):
                            return f"#{rgb}"
                    elif color_obj.type == 'theme':
                        theme_idx = color_obj.theme or 0
                        tint = color_obj.tint or 0
                        base_hex = THEME_COLORS.get(theme_idx, 'FFFFFF')
                        if base_hex in ('FFFFFF', '000000'): return None
                        return f"#{base_hex}"
                except: pass
                return None

            def cell_style(cell, is_header=False):
                styles = []
                has_bg = False
                try:
                    # Fondo
                    if cell.fill and cell.fill.fgColor:
                        bg = get_color_hex(cell.fill.fgColor)
                        if bg and bg.lower() not in ('#000000', '#ffffff', '#000000ff', '#ffffffff'):
                            styles.append(f"background:{bg}")
                            has_bg = True
                    # Fallback gris para headers sin color
                    if is_header and not has_bg:
                        styles.append("background:#f0f0f0")
                        styles.append("font-weight:bold")
                    # Color texto
                    if cell.font and cell.font.color:
                        fg = get_color_hex(cell.font.color)
                        if fg:
                            styles.append(f"color:{fg}")
                    # Negrita (si no está ya en header)
                    if not is_header and cell.font and cell.font.bold:
                        styles.append("font-weight:bold")
                    # Alineación
                    if cell.alignment:
                        if cell.alignment.horizontal == 'center':
                            styles.append("text-align:center")
                        elif cell.alignment.horizontal == 'right':
                            styles.append("text-align:right")
                    # Tamaño fuente
                    if cell.font and cell.font.size:
                        styles.append(f"font-size:{int(cell.font.size)}px")
                except: pass
                return ";".join(styles)

            # Anchos de columna
            col_widths = {}
            for col_letter, col_dim in ws.column_dimensions.items():
                if col_dim.width:
                    col_widths[col_letter] = max(60, int(col_dim.width * 7))

            from openpyxl.utils import get_column_letter

            partes = ['<table style="border-collapse:collapse;font-family:Calibri,Arial,sans-serif;font-size:12px;table-layout:fixed">']

            # Colgroup para anchos + columna de numeración de filas
            partes.append('<colgroup>')
            partes.append('<col style="width:40px;min-width:40px">')  # columna nro fila
            for i in range(1, ws.max_column + 1):
                letter = get_column_letter(i)
                w = col_widths.get(letter, 80)
                partes.append(f'<col style="width:{w}px">')
            partes.append('</colgroup>')

            # Fila de letras de columna
            partes.append('<thead><tr>')
            partes.append('<th style="background:#f2f2f2;border:1px solid #b0b0b0;padding:3px 4px;font-size:11px;color:#666;width:40px;text-align:center"></th>')
            for i in range(1, ws.max_column + 1):
                letter = get_column_letter(i)
                partes.append(f'<th style="background:#f2f2f2;border:1px solid #b0b0b0;padding:3px 6px;font-size:11px;color:#444;text-align:center;font-weight:600">{letter}</th>')
            partes.append('</tr></thead><tbody>')

            row_render = 0
            for row_idx, fila in enumerate(ws.iter_rows()):
                if all(c.value is None for c in fila):
                    continue
                row_render += 1
                is_header = row_idx == 0
                tag_row = 'thead' if is_header else ''
                partes.append('<tr>')
                # Número de fila
                num_style = "background:#f2f2f2;border:1px solid #b0b0b0;padding:3px 4px;font-size:11px;color:#666;text-align:center;font-weight:400;"
                partes.append(f'<td style="{num_style}">{row_idx + 1}</td>')
                for celda in fila:
                    val = celda.value if celda.value is not None else ''
                    # Formatear números con decimales
                    if isinstance(val, float):
                        val = f"{val:g}"
                    style = cell_style(celda, is_header)
                    border = "border:1px solid #d0d0d0;padding:3px 8px;vertical-align:middle;overflow:hidden;text-overflow:ellipsis;"
                    full_style = f"{border}{style}"
                    tag = 'th' if is_header else 'td'
                    partes.append(f'<{tag} style="{full_style}">{val}</{tag}>')
                partes.append('</tr>')

            partes.append('</tbody></table>')
            html = "".join(partes)

        elif ext == "txt":
            texto = contenido_bytes.decode('utf-8', errors='replace')
            lineas = texto.split('\n')
            parrafos = [f"<p>{line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')}</p>" for line in lineas if line.strip()]
            html = "".join(parrafos) or "<p></p>"

        elif ext == "html":
            html = contenido_bytes.decode('utf-8')

        else:
            html = f"<p>Formato .{ext} no soportado para edición en línea. Descarga el archivo para editarlo.</p>"

        return {"contenido_html": html, "nombre": nombre, "formato": ext}

    except Exception as e:
        logger.error(f"Error leyendo documento: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/guardar-contenido",
    summary="Guardar contenido de editor colaborativo",
    description="Guarda contenido HTML/Markdown del editor TipTap en S3/MinIO"
)
async def guardar_contenido_editor(
    key: str = Form(..., description="S3 key del documento"),
    contenido: str = Form(..., description="Contenido HTML o Markdown"),
    formato: str = Form("html", description="Formato: html o markdown")
):
    """Guarda el contenido editado en TipTap al documento en S3/MinIO."""
    if not key:
        raise HTTPException(status_code=400, detail="Key requerida")

    svc = get_s3_service()
    contenido_bytes = contenido.encode('utf-8')

    try:
        svc.subir_contenido_directo(key, contenido_bytes, f'text/{formato}')
        return {"success": True, "key": key, "bytes_guardados": len(contenido_bytes)}
    except Exception as e:
        logger.error(f"Error guardando contenido: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get(
    "/estado/servicio",
    summary="Estado del servicio de almacenamiento",
    description="Verifica si S3/MinIO está disponible y muestra la configuración actual."
)
async def estado_almacenamiento():
    """
    Muestra si el almacenamiento está usando MinIO (local/gratuito) o AWS S3 (producción).
    - MinIO: corre en Docker, mismo API que S3, cero costo
    - AWS S3: producción real en la nube
    """
    svc = get_s3_service()
    return svc.estado()
