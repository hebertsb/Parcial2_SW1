"""
NexusFlow AI — Servicio de Reportes Dinámicos (Segundo Parcial)
===============================================================
Genera reportes en PDF, Excel o Word a partir de una consulta
en lenguaje natural del administrador.

Flujo:
  1. Admin describe reporte (texto libre o filtros estructurados)
  2. Parser extrae: colección, filtros MongoDB, período
  3. Consulta MongoDB con filtro de empresa (Politica.empresa_id → Tramite → Bitacora)
  4. Resuelve IDs → nombres legibles
  5. Enriquece con modelos ML (riesgo_ia, siguiente_accion_predicha)
  6. Genera estadísticas para gráficos + preview de primeras filas
  7. Genera archivo (PDF/Excel/Word) + devuelve todo en response
"""

import base64
import io
import logging
import re
from collections import Counter, defaultdict
from datetime import datetime, timedelta
from typing import Dict, List, Optional

from pymongo import MongoClient

from ..config import settings

logger = logging.getLogger(__name__)

# ── Campos a excluir por colección ───────────────────────────────
_TRAMITE_NOISE = {
    "datos_formulario", "colaboradores", "evidencias", "historial",
    "respuestas_por_nodo", "labels_por_nodo", "_class",
    "cliente_id", "politica_id", "funcionario_asignado_id",  # se resuelven a nombres
    "formulario_actual_id", "nodo_actual_id", "nombre_tramite",
    "tiempo_total", "fecha_ultima_actualizacion",
    "_id",  # excluir ObjectId crudo del output final
}

_BITACORA_NOISE = {
    "usuario_id", "estado_resultante", "_class", "politica_id", "_id",
}

_POLITICA_NOISE = {
    "esquema_workflow",  # dict anidado ilegible
    "_class", "_id",
}

# ── Configuración por colección ───────────────────────────────────
_COL_CONFIG = {
    "Tramite": {
        "fecha_field": "fecha_inicio",
        "estado_field": "estado",
        "cols_priority": [
            "estado", "semaforizacion", "prioridad",
            "cliente_nombre", "politica_nombre",
            "fecha_inicio", "fecha_limite", "fecha_fin",
        ],
    },
    "Bitacora": {
        "fecha_field": "fecha_hora",
        "estado_field": "accion",
        "cols_priority": [
            "tramite_id", "accion", "estado",
            "usuario_nombre", "politica_nombre", "fecha_hora",
        ],
    },
    "Politica": {
        "fecha_field": "fecha_activacion",
        "estado_field": "esta_activa",
        "cols_priority": [
            "nombre", "tipo", "tipo_flujo", "esta_activa",
            "duracion_estandar_dias", "empresa",
            "flujo_nodos", "fecha_activacion",
        ],
    },
}

# ── Colores para gráficos (devueltos al frontend) ─────────────────
_COLORES_ESTADO = {
    "finalizado": "#22c55e", "en_proceso": "#3b82f6",
    "en_revision": "#f59e0b", "rechazado": "#ef4444",
    "pendiente": "#94a3b8", "en_progreso": "#6366f1",
    "observado": "#8b5cf6", "escalado": "#f97316",
    "vencido": "#dc2626",
}

_COLORES_ACCION = {
    "INICIO_PROCESO": "#3b82f6", "LLENADO_FORMULARIO": "#06b6d4",
    "SUBIDA_EVIDENCIA": "#10b981", "APROBAR": "#22c55e",
    "OBSERVAR": "#f59e0b", "RECHAZAR": "#ef4444",
    "ESCALAR": "#f97316", "FINALIZAR": "#8b5cf6",
}


def _get_db():
    client = MongoClient(settings.mongodb_uri, serverSelectionTimeoutMS=5000)
    return client[settings.mongodb_database]


# ── Keywords para parser local ────────────────────────────────────

# Valor especial: tramites activos (no finalizados ni rechazados)
_ESTADOS_ACTIVOS = {"$in": ["en_proceso", "en_progreso", "en_revision", "observado", "escalado"]}

_KEYWORDS_ESTADO = {
    "rechazados": "rechazado", "rechazado": "rechazado",
    "aprobados": "finalizado", "aprobado": "finalizado",
    # "pendiente" no existe como estado en DB — equivale a tramites activos/en curso
    "pendientes": _ESTADOS_ACTIVOS, "pendiente": _ESTADOS_ACTIVOS,
    "activos": _ESTADOS_ACTIVOS, "activo": _ESTADOS_ACTIVOS,
    "en curso": _ESTADOS_ACTIVOS, "en tramite": _ESTADOS_ACTIVOS,
    "sin finalizar": _ESTADOS_ACTIVOS, "no finalizados": _ESTADOS_ACTIVOS,
    "en proceso": "en_proceso", "en_proceso": "en_proceso",
    "en revision": "en_revision", "en_revision": "en_revision",
    "en progreso": "en_progreso", "en_progreso": "en_progreso",
    "completados": "finalizado", "finalizados": "finalizado", "finalizado": "finalizado",
    "observados": "observado", "observado": "observado",
    "vencidos": "vencido", "vencido": "vencido",
    "escalados": "escalado", "escalado": "escalado",
}

_KEYWORDS_COLECCION = {
    "trámites": "Tramite", "tramites": "Tramite", "trámite": "Tramite", "tramite": "Tramite",
    "bitácora": "Bitacora", "bitacora": "Bitacora", "historial": "Bitacora",
    "acciones": "Bitacora", "movimientos": "Bitacora", "log": "Bitacora",
    "recepcionaron": "Bitacora", "recepcionó": "Bitacora", "recepcion": "Bitacora",
    "recepción": "Bitacora", "atendieron": "Bitacora", "atendió": "Bitacora",
    "políticas activas": "Politica", "politicas activas": "Politica",
    "lista de políticas": "Politica", "lista de politicas": "Politica",
}

_KEYWORDS_ACCION = {
    "aprobaciones": "APROBAR", "aprobacion": "APROBAR", "aprobar": "APROBAR", "APROBAR": "APROBAR",
    "rechazos": "RECHAZAR", "rechazo": "RECHAZAR", "rechazar": "RECHAZAR", "RECHAZAR": "RECHAZAR",
    "observaciones": "OBSERVAR", "observacion": "OBSERVAR", "OBSERVAR": "OBSERVAR",
    "escalamientos": "ESCALAR", "escalar": "ESCALAR", "ESCALAR": "ESCALAR",
    "inicios": "INICIO_PROCESO", "inicio de tramite": "INICIO_PROCESO",
    "inicio de trámite": "INICIO_PROCESO", "INICIO_PROCESO": "INICIO_PROCESO",
    "recepcionaron": "INICIO_PROCESO", "recepcionó": "INICIO_PROCESO",
    "finalizaciones": "FINALIZAR", "finalizar": "FINALIZAR", "FINALIZAR": "FINALIZAR",
    "subida": "SUBIDA_EVIDENCIA", "SUBIDA_EVIDENCIA": "SUBIDA_EVIDENCIA",
    "llenado": "LLENADO_FORMULARIO", "LLENADO_FORMULARIO": "LLENADO_FORMULARIO",
}

_KEYWORDS_PERIODO = {
    "este año": 365, "año": 365,
    "últimos 3 meses": 90, "ultimos 3 meses": 90, "3 meses": 90, "tres meses": 90,
    "últimos 2 meses": 60, "ultimos 2 meses": 60,
    "último mes": 30, "ultimo mes": 30, "este mes": 30, "mes": 30,
    "últimas 2 semanas": 14, "ultimas 2 semanas": 14,
    "esta semana": 7, "última semana": 7, "ultima semana": 7, "semana": 7,
    "7 días": 7, "7 dias": 7, "últimos 7 días": 7,
    "últimos 30 días": 30, "ultimos 30 dias": 30,
    "últimos 15 días": 15, "ultimos 15 dias": 15,
    "ayer": 1, "hoy": 0,
}


class ReporteService:
    """Genera reportes dinámicos desde lenguaje natural o filtros estructurados."""

    def generar(
        self,
        consulta: str,
        empresa_id: str,
        formato: str,
        # Filtros estructurados opcionales (desde UI)
        coleccion: Optional[str] = None,
        estado: Optional[str] = None,
        accion: Optional[str] = None,
        fecha_desde: Optional[str] = None,
        fecha_hasta: Optional[str] = None,
        limite: Optional[int] = None,
    ) -> dict:
        formato = formato.lower().strip()
        if formato in ("xlsx",): formato = "excel"
        if formato in ("docx",): formato = "word"
        if formato not in ("pdf", "excel", "word"): formato = "excel"

        # Parser
        if settings.is_api_mode:
            try:
                params = self._parsear_con_ia(consulta, empresa_id)
            except Exception as exc:
                logger.warning("IA no disponible (%s), usando parser local", exc)
                params = self._parsear_local(consulta, empresa_id)
        else:
            params = self._parsear_local(consulta, empresa_id)

        # Aplicar overrides estructurados del frontend
        if coleccion:
            params["coleccion"] = coleccion
            params["tipo"] = "filtro"  # overrides de UI siempre son filtro
        if limite:
            params["limite"] = limite

        cfg = _COL_CONFIG.get(params["coleccion"], _COL_CONFIG["Tramite"])
        fecha_field = cfg["fecha_field"]

        if estado and params["coleccion"] == "Tramite":
            params["filtros"]["estado"] = estado
            # Si estado es finalizado, cualquier filtro de fecha debe ir sobre fecha_fin
            # (cuándo terminó), no sobre fecha_inicio (cuándo empezó)
            if estado == "finalizado":
                for campo_origen in ("fecha_inicio", "fecha_hora"):
                    if campo_origen in params["filtros"] and "fecha_fin" not in params["filtros"]:
                        params["filtros"]["fecha_fin"] = params["filtros"].pop(campo_origen)
        if accion and params["coleccion"] == "Bitacora":
            params["filtros"]["accion"] = accion

        # Rango de fechas explícito (desde frontend o personalizado)
        if fecha_desde or fecha_hasta:
            date_filter: dict = {}
            if fecha_desde:
                date_filter["$gte"] = datetime.fromisoformat(fecha_desde)
            if fecha_hasta:
                hasta_dt = datetime.fromisoformat(fecha_hasta)
                date_filter["$lte"] = hasta_dt.replace(hour=23, minute=59, second=59)
            # Para finalizados el rango aplica sobre fecha_fin
            campo_fecha_destino = "fecha_fin" if (estado == "finalizado") else fecha_field
            params["filtros"][campo_fecha_destino] = date_filter

        # Rutear por tipo de consulta
        tipo_consulta = params.get("tipo", "filtro")

        if tipo_consulta == "ranking":
            datos = self._consultar_ranking(params)
        else:
            # Consulta MongoDB con filtros
            datos = self._consultar_mongo(params)

            # Resolver IDs → nombres legibles
            if params["coleccion"] == "Tramite" and datos:
                datos = self._resolver_tramites(datos)
            elif params["coleccion"] == "Bitacora" and datos:
                datos = self._resolver_bitacora(datos)
            elif params["coleccion"] == "Politica" and datos:
                datos = self._resolver_politica(datos)

            # Enriquecer con predicciones ML
            if params["coleccion"] in ("Tramite", "Bitacora") and datos:
                datos = self._enriquecer_con_ml(datos, params["coleccion"])

        # Calcular columnas, estadísticas y preview para response
        columnas_orden = self._extraer_columnas(datos, params["coleccion"]) if datos else []
        estadisticas = self._generar_estadisticas(datos, params["coleccion"]) if datos else {"total": 0}
        datos_preview = self._serializar_preview(datos, columnas_orden) if datos else []

        logger.info(
            "Reporte: col=%s filtros=%s registros=%d formato=%s",
            params["coleccion"], params.get("filtros", {}), len(datos), formato,
        )

        if formato == "pdf":
            archivo, mime = self._generar_pdf(datos, params, consulta)
            ext = "pdf"
        elif formato == "word":
            archivo, mime = self._generar_word(datos, params, consulta)
            ext = "docx"
        else:
            archivo, mime = self._generar_excel(datos, params, consulta)
            ext = "xlsx"

        nombre = self._nombre_archivo(params, ext)
        return {
            "archivo_b64": base64.b64encode(archivo).decode("utf-8"),
            "nombre_archivo": nombre,
            "mime_type": mime,
            "total_registros": len(datos),
            "consulta_interpretada": params.get("descripcion", consulta),
            "formato": formato,
            "coleccion_consultada": params["coleccion"],
            "datos_preview": datos_preview,
            "estadisticas": estadisticas,
            "columnas": columnas_orden,
        }

    # ── Parsers ───────────────────────────────────────────────────

    def _parsear_local(self, consulta: str, empresa_id: str) -> dict:
        texto = consulta.lower()

        coleccion = "Tramite"
        for kw, col in _KEYWORDS_COLECCION.items():
            if kw in texto:
                coleccion = col
                break

        estado = None
        for kw, est in _KEYWORDS_ESTADO.items():
            if kw in texto:
                estado = est
                break

        accion = None
        if coleccion == "Bitacora":
            for kw, ac in _KEYWORDS_ACCION.items():
                if kw.lower() in texto:
                    accion = ac
                    break

        dias = None
        for kw, d in sorted(_KEYWORDS_PERIODO.items(), key=lambda x: -len(x[0])):
            if kw in texto:
                dias = d
                break

        # Detección libre: "últimos 15 días", "últimas 3 semanas", "últimos 2 meses"
        if dias is None:
            m_dias = re.search(r"últimos?\s+(\d+)\s*d[íi]as?|ultimos?\s+(\d+)\s*dias?", texto)
            m_sem  = re.search(r"últimas?\s+(\d+)\s*semanas?|ultimas?\s+(\d+)\s*semanas?", texto)
            m_mes  = re.search(r"últimos?\s+(\d+)\s*mes(?:es)?|ultimos?\s+(\d+)\s*mes(?:es)?", texto)
            if m_dias:
                dias = int(m_dias.group(1) or m_dias.group(2))
            elif m_sem:
                dias = int(m_sem.group(1) or m_sem.group(2)) * 7
            elif m_mes:
                dias = int(m_mes.group(1) or m_mes.group(2)) * 30

        limite = 500
        m = re.search(r"(\d+)\s*(registros?|tramites?|resultados?|ultimos?|últimos?)", texto)
        if m:
            limite = min(int(m.group(1)), 1000)

        cfg = _COL_CONFIG.get(coleccion, _COL_CONFIG["Tramite"])
        fecha_field = cfg["fecha_field"]
        estado_field = cfg["estado_field"]

        filtros: dict = {}
        if estado and coleccion == "Tramite":
            filtros["estado"] = estado  # puede ser str o dict $in
        if accion and coleccion == "Bitacora":
            filtros["accion"] = accion

        # Para "finalizado" el período aplica sobre fecha_fin (cuándo terminó),
        # no fecha_inicio (cuándo comenzó)
        campo_fecha_query = fecha_field
        if coleccion == "Tramite" and estado == "finalizado" and dias:
            campo_fecha_query = "fecha_fin"

        if dias is not None and dias > 0:
            desde = datetime.now() - timedelta(days=dias)
            filtros[campo_fecha_query] = {"$gte": desde}
        elif dias == 0:
            inicio_hoy = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
            filtros[campo_fecha_query] = {"$gte": inicio_hoy}

        partes = []
        if estado:
            label = "activos/en-curso" if isinstance(estado, dict) else estado
            partes.append(f"estado={label}")
        if accion: partes.append(f"accion={accion}")
        if dias is not None:
            partes.append(f"últimos {dias} días" if dias > 0 else "hoy")
        descripcion = (
            f"{coleccion} "
            + (f"— {', '.join(partes)}" if partes else "— todos los registros")
        )

        # Detección de ranking/agrupación
        _KW_RANKING = [
            "más solicitados", "mas solicitados", "más realizados", "mas realizados",
            "más consultados", "mas consultados", "ranking", "top ", "cuántos por",
            "cuantos por", "más frecuentes", "mas frecuentes", "más usados", "mas usados",
            "estadística", "estadistica", "agrupado", "por tipo", "por estado",
            "qué funcionario", "que funcionario", "qué funcionarios", "que funcionarios",
            "qué cliente", "que cliente", "cuántos trámites", "cuantos tramites",
            "quién atendió", "quien atendio", "funcionarios que más", "funcionarios que mas",
        ]
        es_ranking = any(kw in texto for kw in _KW_RANKING)

        if es_ranking:
            # Determinar campo de agrupación + colección correcta para el conteo
            if any(k in texto for k in ("funcionario", "atendió", "atendio")) and coleccion == "Bitacora":
                # Bitácora: agrupar por usuario_id (quién hizo la acción)
                agrupar_por = "usuario_id"
                col_ranking = "Bitacora"
                descripcion = "Ranking de funcionarios por acciones en bitácora"
            elif any(k in texto for k in ("funcionario", "atendió", "atendio", "recepcionó", "recepcion")):
                agrupar_por = "funcionario_asignado_id"
                col_ranking = "Tramite"
                descripcion = "Ranking de funcionarios por trámites atendidos"
            elif any(k in texto for k in ("cliente",)):
                agrupar_por = "cliente_id"
                col_ranking = "Tramite"
                descripcion = "Ranking de clientes por cantidad de trámites"
            elif any(k in texto for k in ("estado",)):
                agrupar_por = "estado"
                col_ranking = "Tramite"
                descripcion = "Trámites agrupados por estado"
            else:
                # Ranking por política (más solicitadas) — siempre en Tramite
                agrupar_por = "politica_id"
                col_ranking = "Tramite"
                descripcion = "Ranking de políticas/trámites más solicitados"

            return {
                "tipo": "ranking",
                "coleccion": col_ranking,
                "agrupar_por": agrupar_por,
                "filtros": filtros,
                "limite": min(limite, 20),
                "descripcion": descripcion,
                "empresa_id": empresa_id,
            }

        return {
            "tipo": "filtro",
            "coleccion": coleccion,
            "filtros": filtros,
            "limite": limite,
            "descripcion": descripcion,
            "empresa_id": empresa_id,
        }

    def _parsear_con_ia(self, consulta: str, empresa_id: str) -> dict:
        import json
        from openai import OpenAI

        client = OpenAI(
            api_key=settings.openai_api_key,
            base_url=settings.openai_base_url or None,
        )

        system_prompt = (
            "Eres un asistente que convierte consultas en español a parámetros MongoDB para NexusFlow.\n\n"
            "Colecciones disponibles (campos snake_case):\n"
            "  - Tramite: estado (en_proceso|en_progreso|en_revision|finalizado|rechazado|observado|vencido|escalado),\n"
            "    semaforizacion (Verde|Amarillo|Rojo), prioridad (Alta|Media|Baja),\n"
            "    fecha_inicio (Date — cuándo empezó), fecha_fin (Date — cuándo finalizó, SOLO existe si estado=finalizado),\n"
            "    cliente_id, politica_id, funcionario_asignado_id\n"
            "  - Bitacora: accion (INICIO_PROCESO|LLENADO_FORMULARIO|APROBAR|OBSERVAR|RECHAZAR|ESCALAR|SUBIDA_EVIDENCIA|FINALIZAR),\n"
            "    tramite_id, usuario_id, fecha_hora (Date)\n"
            "  - Politica: nombre, tipo_flujo, esta_activa (bool), fecha_activacion\n\n"
            "TIPOS DE CONSULTA:\n"
            "  tipo='filtro'  → consulta con filtros simples, devuelve lista de documentos\n"
            "  tipo='ranking' → agrupación/conteo, devuelve TOP N agrupados por un campo\n\n"
            "CUÁNDO USAR tipo='ranking':\n"
            "  - 'más solicitados', 'más realizados', 'más frecuentes', 'ranking', 'top N'\n"
            "  - 'cuántos trámites por...', 'qué funcionarios atendieron más', 'estadística por estado'\n"
            "  - 'qué funcionario atendió a qué cliente' (agrupar_por=funcionario_asignado_id)\n\n"
            "CAMPO agrupar_por según consulta:\n"
            "  - 'por política/proceso/tipo de trámite' → agrupar_por='politica_id'\n"
            "  - 'por funcionario/quién atendió' → agrupar_por='funcionario_asignado_id'\n"
            "  - 'por cliente/quién solicitó' → agrupar_por='cliente_id'\n"
            "  - 'por estado' → agrupar_por='estado'\n"
            "  - Bitácora por usuario/acción → coleccion='Bitacora', agrupar_por='usuario_id'\n\n"
            "REGLAS CRÍTICAS:\n"
            "  - NO uses empresa_id, empresaId, clienteId, politicaId en filtros.\n"
            "  - Para fechas Bitacora: fecha_hora. Para Tramite inicio: fecha_inicio.\n"
            "  - Si estado=finalizado con período → usar fecha_fin (cuándo terminó), NO fecha_limite.\n"
            "  - $gte/$lte deben ser timestamps ISO8601.\n"
            "  - 'pendiente/pendientes/activos/en curso' = {\"estado\": {\"$in\": [\"en_proceso\",\"en_progreso\",\"en_revision\",\"observado\",\"escalado\"]}}\n"
            "  - 'finalizados/completados' = {\"estado\": \"finalizado\"}\n"
            "  - 'rechazados' = {\"estado\": \"rechazado\"}\n\n"
            "FORMATO DE RESPUESTA (SOLO JSON):\n"
            "  Para tipo='filtro': {\"tipo\":\"filtro\",\"coleccion\":\"Tramite\",\"filtros\":{},\"limite\":200,\"descripcion\":\"...\"}\n"
            "  Para tipo='ranking': {\"tipo\":\"ranking\",\"coleccion\":\"Tramite\",\"agrupar_por\":\"politica_id\",\"filtros\":{},\"limite\":10,\"descripcion\":\"...\"}"
        )

        ahora = datetime.now()
        user_prompt = (
            f'Fecha actual: {ahora.isoformat()}\n'
            f'Consulta: "{consulta}"\n\n'
            'Responde SOLO con JSON válido según el tipo de consulta detectado.'
        )

        response = client.chat.completions.create(
            model=settings.openai_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
        )

        data = json.loads(response.choices[0].message.content)
        data.setdefault("tipo", "filtro")
        data.setdefault("coleccion", "Tramite")
        data.setdefault("filtros", {})
        data.setdefault("limite", 200)
        if data.get("tipo") == "ranking":
            data.setdefault("agrupar_por", "politica_id")
            data.setdefault("limite", 10)
            # Corregir colección: rankings de politica_id/funcionario/cliente → Tramite
            if data.get("agrupar_por") in ("politica_id", "funcionario_asignado_id", "cliente_id"):
                data["coleccion"] = "Tramite"
        data["filtros"].pop("empresa_id", None)
        data["filtros"].pop("empresaId", None)
        data["empresa_id"] = empresa_id
        return data

    # ── Consulta MongoDB ──────────────────────────────────────────

    def _consultar_mongo(self, params: dict) -> list:
        try:
            db = _get_db()
            coleccion = params["coleccion"]
            filtros = dict(params.get("filtros", {}))
            limite = params.get("limite", 200)
            empresa_id = params.get("empresa_id", "")

            cfg = _COL_CONFIG.get(coleccion, _COL_CONFIG["Tramite"])
            fecha_field = cfg["fecha_field"]

            # Constraint de empresa: solo aplica si empresa_id parece un ID real (EMP-xxx)
            empresa_filtro: dict = {}
            if empresa_id and empresa_id.upper().startswith("EMP-"):
                if coleccion == "Politica":
                    empresa_filtro["empresa_id"] = empresa_id
                elif coleccion == "Tramite":
                    pol_ids = [str(p) for p in db.Politica.distinct("_id", {"empresa_id": empresa_id})]
                    if pol_ids:
                        empresa_filtro["politica_id"] = {"$in": pol_ids}
                elif coleccion == "Bitacora":
                    # Si hay más de una empresa registrada, aislar Bitácora por los
                    # trámites de esa empresa. Con una sola empresa, toda la Bitácora
                    # (incluye histórico/dataset de entrenamiento) le pertenece.
                    if len(db.Politica.distinct("empresa_id")) > 1:
                        pol_ids = [str(p) for p in db.Politica.distinct("_id", {"empresa_id": empresa_id})]
                        if pol_ids:
                            t_ids = [str(t) for t in db.Tramite.distinct("_id", {"politica_id": {"$in": pol_ids}})]
                            if t_ids:
                                empresa_filtro["tramite_id"] = {"$in": t_ids}

            filtros_full = {**filtros, **empresa_filtro}
            nombres_candidatos = [coleccion, coleccion.lower(), coleccion + "s", coleccion.lower() + "s"]

            for nombre in nombres_candidatos:
                docs = list(db[nombre].find(filtros_full).limit(limite))
                if docs:
                    logger.info("Consulta OK: col=%s empresa=%s docs=%d", nombre, empresa_id, len(docs))
                    return docs

            # Fallback 1: quitar CUALQUIER filtro de rango fecha ($gte/$lte/$gt/$lt),
            # sin importar el nombre del campo — conserva empresa + estado/accion
            _ops_fecha = {"$gte", "$lte", "$gt", "$lt"}
            filtros_sin_fecha = {
                k: v for k, v in filtros_full.items()
                if not (isinstance(v, dict) and _ops_fecha & set(v.keys()))
            }
            if filtros_sin_fecha != filtros_full:
                for nombre in nombres_candidatos:
                    docs = list(db[nombre].find(filtros_sin_fecha, {"_id": 0}).limit(limite))
                    if docs:
                        logger.info("Fallback sin fecha: col=%s docs=%d", nombre, len(docs))
                        params["descripcion"] = params.get("descripcion", "") + " (histórico — sin filtro de fecha)"
                        return docs

            # Fallback 2: solo filtros de usuario (estado/accion), sin empresa ni fecha
            # Preserva la intención del usuario aunque no haya datos de su empresa
            filtros_usuario = {k: v for k, v in filtros.items() if k != fecha_field}
            if filtros_usuario:
                for nombre in nombres_candidatos:
                    docs = list(db[nombre].find(filtros_usuario, {"_id": 0}).limit(limite))
                    if docs:
                        logger.info("Fallback filtros-usuario: col=%s docs=%d", nombre, len(docs))
                        params["descripcion"] = params.get("descripcion", "") + " (histórico — todos los períodos)"
                        return docs

            # Fallback 3: sin ningún filtro — SOLO para super admin (sin empresa real)
            if not empresa_id.upper().startswith("EMP-"):
                for nombre in nombres_candidatos:
                    docs = list(db[nombre].find({}, {"_id": 0}).limit(limite))
                    if docs:
                        logger.info("Fallback sin filtros: col=%s docs=%d", nombre, len(docs))
                        params["descripcion"] = params.get("descripcion", "") + " (todos los registros)"
                        return docs

            logger.warning("Sin datos en '%s' empresa=%s", coleccion, empresa_id)
            return []
        except Exception as exc:
            logger.error("Error MongoDB: %s", exc)
            return []

    # ── Consulta de ranking / agregación ─────────────────────────

    def _consultar_ranking(self, params: dict) -> list:
        """Ejecuta un pipeline de agregación MongoDB para queries de tipo
        'más solicitados', 'top políticas', 'funcionarios con más acciones', etc."""
        try:
            from bson import ObjectId
            db = _get_db()
            coleccion   = params.get("coleccion", "Tramite")
            agrupar_por = params.get("agrupar_por", "politica_id")
            empresa_id  = params.get("empresa_id", "")
            limite      = params.get("limite", 10)
            filtros     = dict(params.get("filtros", {}))

            # Empresa constraint
            empresa_filtro: dict = {}
            if empresa_id and empresa_id.upper().startswith("EMP-"):
                if coleccion == "Tramite":
                    pol_ids = [str(p) for p in db.Politica.distinct("_id", {"empresa_id": empresa_id})]
                    if pol_ids:
                        empresa_filtro["politica_id"] = {"$in": pol_ids}
                elif coleccion == "Bitacora":
                    # Si hay más de una empresa registrada, aislar Bitácora por los
                    # trámites de esa empresa. Con una sola empresa, toda la Bitácora
                    # (incluye histórico/dataset de entrenamiento) le pertenece.
                    if len(db.Politica.distinct("empresa_id")) > 1:
                        pol_ids = [str(p) for p in db.Politica.distinct("_id", {"empresa_id": empresa_id})]
                        if pol_ids:
                            t_ids = [str(t) for t in db.Tramite.distinct("_id", {"politica_id": {"$in": pol_ids}})]
                            if t_ids:
                                empresa_filtro["tramite_id"] = {"$in": t_ids}

            match_stage = {**filtros, **empresa_filtro}

            pipeline = [
                {"$match": match_stage} if match_stage else {"$match": {}},
                {"$group": {
                    "_id": f"${agrupar_por}",
                    "total": {"$sum": 1},
                    "ultimo": {"$max": "$fecha_inicio" if coleccion == "Tramite" else "$fecha_hora"},
                }},
                {"$sort": {"total": -1}},
                {"$limit": limite},
            ]

            raw = list(db[coleccion].aggregate(pipeline))
            if not raw:
                return []

            # Resolver el campo agrupado a nombre legible
            ids_agrupados = [str(r["_id"]) for r in raw if r.get("_id")]

            def _nombre_usuario(uid: str) -> str:
                try:
                    u = db.Usuario.find_one({"_id": ObjectId(uid)}, {"nombre_completo": 1})
                    return u.get("nombre_completo", uid) if u else uid
                except Exception:
                    return uid

            label_map: dict = {}
            es_usuario = any(k in agrupar_por for k in ("funcionario", "cliente_id", "usuario_id"))

            if "politica" in agrupar_por:
                campo_label = "politica_nombre"
                for pid in ids_agrupados:
                    try:
                        p = db.Politica.find_one({"_id": ObjectId(pid)}, {"nombre": 1})
                    except Exception:
                        p = None
                    label_map[pid] = p.get("nombre", pid) if p else pid

            elif es_usuario:
                # funcionario_asignado_id, cliente_id, usuario_id — todos son ObjectId → Usuario
                if "funcionario" in agrupar_por:
                    campo_label = "funcionario_nombre"
                elif "cliente" in agrupar_por:
                    campo_label = "cliente_nombre"
                else:
                    campo_label = "usuario_nombre"
                for uid in ids_agrupados:
                    label_map[uid] = _nombre_usuario(uid)

            elif "estado" in agrupar_por:
                campo_label = "estado"
                label_map = {r: r for r in ids_agrupados}

            else:
                campo_label = agrupar_por
                label_map = {r: r for r in ids_agrupados}

            result = []
            for i, r in enumerate(raw, 1):
                raw_id = str(r.get("_id", ""))
                nombre_resuelto = label_map.get(raw_id, raw_id or "—")
                result.append({
                    "posicion": i,
                    campo_label: nombre_resuelto,
                    "total_tramites": r["total"],
                    "ultima_actividad": r.get("ultimo", "—"),
                })
            return result
        except Exception as exc:
            logger.error("_consultar_ranking falló: %s", exc)
            return []

    # ── Resolución de IDs → nombres ───────────────────────────────

    def _resolver_tramites(self, docs: list) -> list:
        """Reemplaza IDs por nombres legibles: cliente, política, funcionario asignado,
        y último funcionario que actuó (desde Bitácora)."""
        try:
            from bson import ObjectId
            db = _get_db()

            cliente_ids, politica_ids, func_ids, tramite_ids = set(), set(), set(), set()
            for doc in docs:
                if doc.get("cliente_id"):             cliente_ids.add(str(doc["cliente_id"]))
                if doc.get("politica_id"):            politica_ids.add(str(doc["politica_id"]))
                if doc.get("funcionario_asignado_id"):func_ids.add(str(doc["funcionario_asignado_id"]))
                raw_id = doc.get("_id")
                if raw_id: tramite_ids.add(str(raw_id))

            def _nombre_usuario(uid: str) -> str:
                try:
                    u = db.Usuario.find_one({"_id": ObjectId(uid)}, {"nombre_completo": 1})
                    return u.get("nombre_completo", uid) if u else uid
                except Exception:
                    return uid

            def _nombre_politica(pid: str) -> str:
                try:
                    p = db.Politica.find_one({"_id": ObjectId(pid)}, {"nombre": 1})
                except Exception:
                    p = None
                if not p:
                    p = db.Politica.find_one({"_id": pid}, {"nombre": 1})
                return p.get("nombre", pid) if p else pid

            cliente_map  = {cid: _nombre_usuario(cid)   for cid in cliente_ids}
            func_map     = {fid: _nombre_usuario(fid)   for fid in func_ids}
            politica_map = {pid: _nombre_politica(pid)  for pid in politica_ids}

            # Último funcionario que actuó en cada trámite (última entrada Bitácora)
            ultimo_func_map: dict = {}
            if tramite_ids:
                for entrada in db.Bitacora.find(
                    {"tramite_id": {"$in": list(tramite_ids)}},
                    {"tramite_id": 1, "usuario_id": 1, "accion": 1, "fecha_hora": 1}
                ).sort("fecha_hora", -1):
                    tid = str(entrada.get("tramite_id", ""))
                    if tid and tid not in ultimo_func_map:
                        uid = str(entrada.get("usuario_id", ""))
                        accion = entrada.get("accion", "")
                        nombre = _nombre_usuario(uid) if uid else "—"
                        ultimo_func_map[tid] = f"{nombre} ({accion})" if accion else nombre

            result = []
            for doc in docs:
                d = {k: v for k, v in doc.items() if k not in _TRAMITE_NOISE}
                cid = str(doc.get("cliente_id", ""))
                pid = str(doc.get("politica_id", ""))
                fid = str(doc.get("funcionario_asignado_id", ""))
                tid = str(doc.get("_id", ""))
                d["cliente_nombre"]    = cliente_map.get(cid, "—")
                d["politica_nombre"]   = politica_map.get(pid, "—")
                d["funcionario_nombre"]= func_map.get(fid, "—") if fid else "—"
                d["ultimo_funcionario"]= ultimo_func_map.get(tid, "—")
                result.append(d)
            return result
        except Exception as exc:
            logger.warning("_resolver_tramites falló: %s", exc)
            return docs

    def _resolver_bitacora(self, docs: list) -> list:
        """Resuelve usuario_id → nombre_completo, politica_id → nombre, tramite_id → nombre legible."""
        try:
            from bson import ObjectId
            db = _get_db()

            usuario_ids, politica_ids, tramite_ids = set(), set(), set()
            for doc in docs:
                if doc.get("usuario_id"):  usuario_ids.add(str(doc["usuario_id"]))
                if doc.get("politica_id"): politica_ids.add(str(doc["politica_id"]))
                if doc.get("tramite_id"):  tramite_ids.add(str(doc["tramite_id"]))

            usuario_map: dict = {}
            for uid in usuario_ids:
                try:
                    u = db.Usuario.find_one({"_id": ObjectId(uid)}, {"nombre_completo": 1})
                    usuario_map[uid] = u.get("nombre_completo", uid) if u else uid
                except Exception:
                    usuario_map[uid] = uid

            politica_map: dict = {}
            for pid in politica_ids:
                try:
                    p = db.Politica.find_one({"_id": ObjectId(pid)}, {"nombre": 1})
                except Exception:
                    p = None
                if not p:
                    p = db.Politica.find_one({"_id": pid}, {"nombre": 1})
                politica_map[pid] = p.get("nombre", pid) if p else pid

            # Resolver tramite_id → nombre de la política del trámite + ID corto
            tramite_label_map: dict = {}
            tramite_politica_map: dict = {}
            for tid in tramite_ids:
                short = f"#{tid[:6]}"
                try:
                    t = db.Tramite.find_one({"_id": ObjectId(tid)}, {"politica_id": 1, "nombre_tramite": 1})
                    if t:
                        pol_id = str(t.get("politica_id", ""))
                        nombre_politica = politica_map.get(pol_id, "")
                        if not nombre_politica and pol_id:
                            try:
                                p2 = db.Politica.find_one({"_id": ObjectId(pol_id)}, {"nombre": 1})
                                nombre_politica = p2.get("nombre", "") if p2 else ""
                            except Exception:
                                nombre_politica = ""
                        tramite_politica_map[tid] = nombre_politica or "—"

                        nombre = t.get("nombre_tramite") or nombre_politica
                        tramite_label_map[tid] = f"{nombre} {short}" if nombre else short
                    else:
                        tramite_label_map[tid] = short
                        tramite_politica_map[tid] = "—"
                except Exception:
                    tramite_label_map[tid] = short
                    tramite_politica_map[tid] = "—"

            result = []
            for doc in docs:
                d = {k: v for k, v in doc.items() if k not in _BITACORA_NOISE}
                uid = str(doc.get("usuario_id", ""))
                pid = str(doc.get("politica_id", ""))
                tid = str(doc.get("tramite_id", ""))
                d["usuario_nombre"]  = usuario_map.get(uid, "—")
                d["politica_nombre"] = politica_map.get(pid) or tramite_politica_map.get(tid, "—")
                if tid:
                    d["tramite_id"] = tramite_label_map.get(tid, f"#{tid[:6]}")
                # Simplificar detalle_ia
                det = d.get("detalle_ia")
                if isinstance(det, dict):
                    partes = []
                    if "tiempo_en_nodo" in det:
                        partes.append(f"t={det['tiempo_en_nodo']}h")
                    if "campos_configurados" in det:
                        partes.append(f"campos={det['campos_configurados']}")
                    if "numero_documentos" in det:
                        partes.append(f"docs={det['numero_documentos']}")
                    d["detalle_ia"] = " | ".join(partes) if partes else "—"
                result.append(d)
            return result
        except Exception as exc:
            logger.warning("_resolver_bitacora falló: %s", exc)
            return docs

    def _resolver_politica(self, docs: list) -> list:
        """Limpia Política: resuelve empresa_id → nombre, añade flujo_nodos legible."""
        try:
            db = _get_db()

            empresa_ids = {str(d.get("empresa_id", "")) for d in docs if d.get("empresa_id")}
            _EMPRESA_FALLBACK = {"EMP-DEFAULT": "Empresa por Defecto"}
            empresa_map: dict = {}
            for eid in empresa_ids:
                e = db.Empresa.find_one({"_id": eid}, {"nombre_legal": 1, "nombre": 1, "razon_social": 1})
                if not e:
                    e = db.Empresa.find_one({"empresa_id": eid}, {"nombre_legal": 1, "nombre": 1})
                if e:
                    empresa_map[eid] = (
                        e.get("nombre_legal") or e.get("nombre") or
                        e.get("razon_social") or eid
                    )
                else:
                    empresa_map[eid] = _EMPRESA_FALLBACK.get(eid, eid)

            result = []
            for doc in docs:
                d = {k: v for k, v in doc.items() if k not in _POLITICA_NOISE}
                # Empresa
                eid = str(d.pop("empresa_id", "") or "")
                d["empresa"] = empresa_map.get(eid, eid) if eid else "—"
                # Nodos del flujo
                workflow = doc.get("esquema_workflow")
                d["flujo_nodos"] = _extraer_nodos_workflow(workflow) if isinstance(workflow, dict) else "—"
                result.append(d)
            return result
        except Exception as exc:
            logger.warning("_resolver_politica falló: %s", exc)
            return [{k: v for k, v in d.items() if k not in _POLITICA_NOISE} for d in docs]

    # ── Estadísticas para gráficos ────────────────────────────────

    def _generar_estadisticas(self, datos: list, coleccion: str) -> dict:
        stats: dict = {"total": len(datos)}

        if coleccion == "Tramite":
            stats["por_estado"] = dict(Counter(d.get("estado", "—") for d in datos))
            stats["por_semaforo"] = dict(Counter(d.get("semaforizacion", "—") for d in datos))
            stats["colores_estado"] = {k: _COLORES_ESTADO.get(k, "#94a3b8") for k in stats["por_estado"]}
            timeline: dict = defaultdict(int)
            for d in datos:
                f = d.get("fecha_inicio")
                if isinstance(f, datetime):
                    timeline[f.strftime("%Y-%m")] += 1
            stats["timeline"] = dict(sorted(timeline.items()))
            stats["timeline_label"] = "por mes"

        elif coleccion == "Bitacora":
            stats["por_accion"] = dict(Counter(d.get("accion", "—") for d in datos))
            stats["por_estado"] = dict(Counter(d.get("estado", "—") for d in datos))
            stats["colores_accion"] = {k: _COLORES_ACCION.get(k, "#94a3b8") for k in stats["por_accion"]}
            timeline2: dict = defaultdict(int)
            for d in datos:
                f = d.get("fecha_hora")
                if isinstance(f, datetime):
                    timeline2[f.strftime("%d/%m")] += 1
            sorted_tl = dict(sorted(timeline2.items()))
            # Últimas 20 entradas
            keys = list(sorted_tl.keys())[-20:]
            stats["timeline"] = {k: sorted_tl[k] for k in keys}
            stats["timeline_label"] = "por día"

        elif coleccion == "Politica":
            activas = sum(1 for d in datos if d.get("activa") is True)
            stats["por_estado"] = {"activa": activas, "inactiva": len(datos) - activas}
            stats["colores_estado"] = {"activa": "#22c55e", "inactiva": "#94a3b8"}

        return stats

    # ── Preview JSON para tabla en browser ────────────────────────

    def _serializar_preview(self, datos: list, columnas: list, n: int = 20) -> list:
        result = []
        for doc in datos[:n]:
            row: dict = {}
            for col in columnas:
                val = doc.get(col, "")
                val = _formatear_valor(val)
                if isinstance(val, (dict, list)):
                    val = str(val)[:80]
                row[col] = val if val is not None else ""
            result.append(row)
        return result

    # ── Enriquecimiento ML ────────────────────────────────────────

    def _enriquecer_con_ml(self, datos: list, coleccion: str) -> list:
        try:
            from .tensorflow_service import get_tf_service
            tf_svc = get_tf_service()
            ahora = datetime.now()

            enriquecidos = []
            for doc in datos:
                doc = dict(doc)

                if coleccion == "Tramite":
                    estado = doc.get("estado", "")
                    fecha_creacion = doc.get("fecha_inicio") or doc.get("fechaCreacion")
                    horas_transcurridas = 0.0
                    if isinstance(fecha_creacion, datetime):
                        horas_transcurridas = (ahora - fecha_creacion).total_seconds() / 3600

                    features = {
                        "tiempo_en_nodo": horas_transcurridas,
                        "campos_configurados": 5.0,
                        "numero_documentos": 2.0,
                        "hora_dia": float(ahora.hour),
                        "dia_semana": float(ahora.weekday()),
                        "ultima_accion": _estado_a_accion(estado),
                    }
                    resultado_riesgo = tf_svc.predecir_riesgo_demora(features)
                    doc["riesgo_ia"] = resultado_riesgo.get("nivel_riesgo", "—")
                    doc["riesgo_probabilidad"] = resultado_riesgo.get("probabilidad_demora", 0.0)

                elif coleccion == "Bitacora":
                    accion = doc.get("accion", "")
                    pol_id = doc.get("politica_id")
                    resultado_lstm = tf_svc.predecir_ruta(
                        [accion],
                        politica_id=pol_id if isinstance(pol_id, str) else None,
                    )
                    doc["siguiente_accion_predicha"] = resultado_lstm.get("siguiente_accion", "—")
                    doc["confianza_prediccion"] = resultado_lstm.get("confianza", 0.0)

                enriquecidos.append(doc)
            return enriquecidos
        except Exception as exc:
            logger.warning("ML enrichment falló (%s) — sin predicciones", exc)
            return datos

    # ── Generadores de archivo ────────────────────────────────────

    def _generar_excel(self, datos: list, params: dict, consulta: str) -> tuple:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Reporte NexusFlow"

        ws.merge_cells("A1:J1")
        ws["A1"] = "NexusFlow — Reporte Dinámico"
        ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
        ws["A1"].fill = PatternFill("solid", fgColor="1E3A5F")
        ws["A1"].alignment = Alignment(horizontal="center")

        ws.merge_cells("A2:J2")
        ws["A2"] = f"Consulta: {consulta}"
        ws["A2"].font = Font(italic=True, size=10)

        ws.merge_cells("A3:J3")
        ws["A3"] = (
            f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')} | "
            f"Total: {len(datos)} registros | Colección: {params['coleccion']}"
        )
        ws["A3"].font = Font(size=9, color="666666")

        if not datos:
            ws["A5"] = "No se encontraron registros con los criterios especificados."
            ws["A5"].font = Font(italic=True, color="888888")
        else:
            columnas = self._extraer_columnas(datos, params["coleccion"])
            header_row = 5

            for col_idx, col_name in enumerate(columnas, 1):
                cell = ws.cell(row=header_row, column=col_idx, value=col_name.upper().replace("_", " "))
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="2563EB")
                cell.alignment = Alignment(horizontal="center", wrap_text=True)

            for row_idx, doc in enumerate(datos, header_row + 1):
                fill = PatternFill("solid", fgColor="F0F4FF") if row_idx % 2 == 0 else None
                for col_idx, col_name in enumerate(columnas, 1):
                    val = doc.get(col_name, "")
                    val = _formatear_valor(val)
                    cell = ws.cell(row=row_idx, column=col_idx, value=val)
                    cell.alignment = Alignment(wrap_text=False)
                    if fill:
                        cell.fill = fill
                    if col_name == "riesgo_ia" and isinstance(val, str):
                        color = {"CRÍTICO": "FF4444", "ALTO": "FF8C00",
                                 "MEDIO": "FFD700", "BAJO": "22C55E"}.get(val)
                        if color:
                            cell.font = Font(color=color, bold=True)
                    if col_name == "estado" and isinstance(val, str):
                        color_e = {
                            "finalizado": "22C55E", "rechazado": "EF4444",
                            "en_proceso": "3B82F6", "en_revision": "F59E0B",
                            "observado": "8B5CF6", "escalado": "F97316",
                        }.get(val)
                        if color_e:
                            cell.font = Font(color=color_e, bold=True)

            # Anchos automáticos
            col_min_widths = {
                "estado": 14, "semaforizacion": 14, "prioridad": 12,
                "cliente_nombre": 22, "politica_nombre": 26,
                "nodo_actual_id": 14, "fecha_inicio": 18, "fecha_limite": 18,
                "riesgo_ia": 14, "siguiente_accion_predicha": 26,
                "accion": 22, "tramite_id": 26, "usuario_nombre": 22,
                "politica_nombre": 28, "fecha_hora": 18, "usuario_id": 26,
            }
            for col_idx2, col_name2 in enumerate(columnas, 1):
                col_letter = ws.cell(row=header_row, column=col_idx2).column_letter
                data_vals = [
                    ws.cell(row=r, column=col_idx2).value
                    for r in range(header_row, header_row + min(len(datos), 50) + 1)
                ]
                max_data = max((len(str(v or "")) for v in data_vals), default=8)
                min_w = col_min_widths.get(col_name2, 14)
                ws.column_dimensions[col_letter].width = min(max(max_data + 2, min_w), 36)

            ws.row_dimensions[header_row].height = 30

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    def _generar_pdf(self, datos: list, params: dict, consulta: str) -> tuple:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

        # Anchos preferidos por columna (en cm)
        _W = {
            "estado": 2.0, "accion": 2.2, "semaforizacion": 1.8, "prioridad": 1.8,
            "esta_activa": 1.8, "activa": 1.8, "tipo_flujo": 2.0, "tipo": 2.0,
            "duracion_estandar_dias": 2.2, "empresa": 2.2, "flujo_nodos": 4.5,
            "cliente_nombre": 2.8, "politica_nombre": 3.0, "usuario_nombre": 2.8,
            "nombre": 3.5,
            "fecha_inicio": 2.3, "fecha_limite": 2.3, "fecha_fin": 2.3,
            "fecha_hora": 2.3, "fecha_activacion": 2.3,
            "tramite_id": 2.8,
            "riesgo_ia": 1.8, "siguiente_accion_predicha": 2.5,
        }

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                                leftMargin=1.5*cm, rightMargin=1.5*cm,
                                topMargin=1.5*cm, bottomMargin=1.5*cm)
        styles = getSampleStyleSheet()
        story = []

        cell_style = ParagraphStyle("cell", fontSize=7, leading=9,
                                    fontName="Helvetica", wordWrap="LTR")
        hdr_style  = ParagraphStyle("hdr",  fontSize=7, leading=9,
                                    fontName="Helvetica-Bold",
                                    textColor=colors.white, wordWrap="LTR")
        title_style = ParagraphStyle("title", parent=styles["Title"],
                                     textColor=colors.HexColor("#1E3A5F"), fontSize=16)
        meta_style  = ParagraphStyle("meta",  parent=styles["Normal"],
                                     fontSize=9, textColor=colors.grey)

        story.append(Paragraph("NexusFlow — Reporte Dinámico", title_style))
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(f"<i>Consulta: {consulta[:120]}</i>", styles["Normal"]))
        story.append(Paragraph(
            f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')} | "
            f"Total: {len(datos)} registros | Colección: {params['coleccion']}",
            meta_style,
        ))
        story.append(Spacer(1, 0.5*cm))

        if not datos:
            story.append(Paragraph("No se encontraron registros.", styles["Normal"]))
        else:
            columnas = self._extraer_columnas(datos, params["coleccion"])
            header = [[Paragraph(c.upper().replace("_", " "), hdr_style)] for c in columnas]
            tabla_data = [header]

            for doc_row in datos:
                fila = []
                for col in columnas:
                    val = _formatear_valor(doc_row.get(col, ""))
                    txt = str(val)[:60] if val else "—"
                    fila.append(Paragraph(txt, cell_style))
                tabla_data.append(fila)

            page_w = landscape(A4)[0] - 3*cm
            raw_widths = [_W.get(c, 2.2) * cm for c in columnas]
            total_w = sum(raw_widths)
            if total_w > page_w:
                scale = page_w / total_w
                col_widths = [w * scale for w in raw_widths]
            else:
                col_widths = raw_widths

            table = Table(tabla_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#2563EB")),
                ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4FF")]),
                ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("ALIGN",         (0, 0), (-1, -1), "LEFT"),
                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING",    (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING",   (0, 0), (-1, -1), 4),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
            ]))
            story.append(table)

        doc.build(story)
        return buf.getvalue(), "application/pdf"

    def _generar_word(self, datos: list, params: dict, consulta: str) -> tuple:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Anchos preferidos por columna (cm)
        _W = {
            "estado": 2.2, "accion": 2.4, "semaforizacion": 2.0, "prioridad": 2.0,
            "esta_activa": 2.0, "activa": 2.0, "tipo_flujo": 2.2, "tipo": 2.2,
            "duracion_estandar_dias": 2.4, "empresa": 2.4, "flujo_nodos": 5.0,
            "cliente_nombre": 3.2, "politica_nombre": 3.5, "usuario_nombre": 3.2,
            "nombre": 4.0,
            "fecha_inicio": 2.6, "fecha_limite": 2.6, "fecha_fin": 2.6,
            "fecha_hora": 2.6, "fecha_activacion": 2.6,
            "tramite_id": 3.2,
            "riesgo_ia": 2.0, "siguiente_accion_predicha": 3.0,
        }

        doc = Document()

        # Landscape A4
        section = doc.sections[-1]
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin  = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        titulo = doc.add_heading("NexusFlow — Reporte Dinámico", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_consulta = doc.add_paragraph()
        p_consulta.add_run(f"Consulta: {consulta}").italic = True
        meta = doc.add_paragraph(
            f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')} | "
            f"Total: {len(datos)} registros | Colección: {params['coleccion']}"
        )
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.add_paragraph()

        if not datos:
            doc.add_paragraph("No se encontraron registros.")
        else:
            columnas = self._extraer_columnas(datos, params["coleccion"])
            tabla = doc.add_table(rows=1, cols=len(columnas))
            tabla.style = "Table Grid"

            # Header row
            hdr = tabla.rows[0].cells
            for i, col in enumerate(columnas):
                p = hdr[i].paragraphs[0]
                run = p.add_run(col.upper().replace("_", " "))
                run.bold = True
                run.font.size = Pt(8)
                # Header background azul
                tc = hdr[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '2563EB')
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            # Data rows
            for ri, doc_row in enumerate(datos):
                row = tabla.add_row().cells
                fill = "F0F4FF" if ri % 2 == 0 else "FFFFFF"
                for i, col in enumerate(columnas):
                    val = _formatear_valor(doc_row.get(col, ""))
                    txt = str(val)[:80] if val else "—"
                    p = row[i].paragraphs[0]
                    run = p.add_run(txt)
                    run.font.size = Pt(8)
                    # Row alternating fill
                    tc = row[i]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), fill)
                    tcPr.append(shd)

            # Set column widths
            usable_w = 29.7 - 3.0  # landscape A4 minus margins (cm)
            raw = [_W.get(c, 2.4) for c in columnas]
            total_w = sum(raw)
            scale = (usable_w / total_w) if total_w > usable_w else 1.0
            for row_obj in tabla.rows:
                for ci, cell in enumerate(row_obj.cells):
                    cell.width = Cm(raw[ci] * scale)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # ── Utilidades ────────────────────────────────────────────────

    def _extraer_columnas(self, datos: list, coleccion: str = "Tramite") -> list:
        todos = set()
        for doc in datos[:20]:
            todos.update(doc.keys())

        cfg = _COL_CONFIG.get(coleccion, _COL_CONFIG["Tramite"])
        prioritarias = cfg["cols_priority"]
        ml_cols = ["riesgo_ia", "riesgo_probabilidad", "siguiente_accion_predicha", "confianza_prediccion"]

        noise_extra = _TRAMITE_NOISE | _BITACORA_NOISE | _POLITICA_NOISE
        columnas = [c for c in prioritarias if c in todos]
        ml_presentes = [c for c in ml_cols if c in todos]
        excluir = set(ml_cols) | noise_extra | {
            "_id", "_class", "usuario_id", "politica_id", "cliente_id",
            "datos_formulario", "colaboradores", "evidencias", "historial",
            "respuestas_por_nodo", "labels_por_nodo", "estado_resultante",
            "formulario_actual_id", "funcionario_asignado_id",
            "nodo_actual_id", "nombre_tramite", "tiempo_total",
            "fecha_ultima_actualizacion", "esquema_workflow", "empresa_id",
        }
        resto = sorted(todos - set(columnas) - excluir)
        return (columnas + resto + ml_presentes)[:12]

    def _nombre_archivo(self, params: dict, ext: str) -> str:
        fecha = datetime.now().strftime("%Y%m%d_%H%M")
        col = params.get("coleccion", "reporte").lower()
        return f"nexusflow_reporte_{col}_{fecha}.{ext}"


# ── Helpers globales ──────────────────────────────────────────────

_UUID_RE = re.compile(
    r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', re.I
)
_OBJECTID_RE = re.compile(r'^[0-9a-f]{24}$', re.I)


def _es_id_crudo(val: str) -> bool:
    """True si el valor parece un UUID o MongoDB ObjectId sin procesar."""
    return bool(_UUID_RE.match(val) or _OBJECTID_RE.match(val))


def _formatear_valor(val) -> str | float | int:
    if val is None:
        return "—"
    if isinstance(val, bool):
        return "Sí" if val else "No"
    if isinstance(val, datetime):
        return val.strftime("%d/%m/%Y %H:%M")
    if isinstance(val, (dict, list)):
        return "—"
    if isinstance(val, float):
        return round(val, 4)
    if isinstance(val, str):
        stripped = val.strip()
        if _es_id_crudo(stripped):
            return f"#{stripped[:8]}"
    return val


def _extraer_nodos_workflow(workflow: dict) -> str:
    """Convierte esquema_workflow → cadena legible de nodos ordenados."""
    if not isinstance(workflow, dict):
        return "—"

    nodos     = workflow.get("nodos") or workflow.get("nodes") or []
    relaciones = workflow.get("relaciones") or workflow.get("edges") or []

    # ── Construir mapa id→label desde la lista de nodos ──────────
    node_label: dict = {}
    _TIPO_GENERICO = {
        "inicio": "Inicio", "start": "Inicio",
        "fin": "Fin", "end": "Fin", "final": "Fin",
        "decision": "Decisión", "gateway": "Decisión",
        "revision": "Revisión", "review": "Revisión",
        "tarea": "Tarea", "task": "Tarea",
        "aprobacion": "Aprobación", "approval": "Aprobación",
    }
    for i, n in enumerate(nodos, 1):
        nid = str(n.get("id") or n.get("nodeId") or f"n{i}")
        label = (
            (n.get("label") or n.get("nombre") or n.get("name") or
             n.get("titulo") or n.get("title") or "")
        ).strip()
        if not label:
            tipo = str(n.get("tipo") or n.get("type") or "").lower().strip()
            label = _TIPO_GENERICO.get(tipo) or f"Nodo {i}"
        node_label[nid] = label

    # ── Reconstruir cadena desde relaciones ───────────────────────
    if relaciones:
        # Verificar si son relaciones inter-política (tienen 'id' UUID largo) → ignorar
        primer = relaciones[0] if relaciones else {}
        es_interpolitica = (
            "politicaId" in primer or
            (_es_id_crudo(str(primer.get("id", ""))) and "origen" not in primer)
        )
        if es_interpolitica:
            # Solo mostrar tipo de flujo genérico
            tipo_flujo = str(primer.get("tipo") or "").strip()
            if tipo_flujo:
                return f"Flujo {tipo_flujo} ({len(relaciones)} pasos)"
            return f"{len(relaciones)} relaciones inter-política"

        # Relaciones con origen/destino
        adj: dict = {}
        todos_nodos: list = []
        todos_set: set = set()
        for rel in relaciones:
            orig = str(rel.get("origen") or rel.get("source") or rel.get("from") or "").strip()
            dest = str(rel.get("destino") or rel.get("target") or rel.get("to") or "").strip()
            if orig and dest:
                adj.setdefault(orig, []).append(dest)
                for nid in (orig, dest):
                    if nid not in todos_set:
                        todos_nodos.append(nid)
                        todos_set.add(nid)

        if not todos_nodos:
            return "—"

        # Si no tenemos mapa de labels, generar genérico por posición
        if not node_label:
            for i, nid in enumerate(todos_nodos):
                if i == 0:
                    node_label[nid] = "Inicio"
                elif i == len(todos_nodos) - 1:
                    node_label[nid] = "Fin"
                else:
                    node_label[nid] = f"Nodo {i}"

        # BFS desde el nodo raíz (sin entrantes)
        entrantes = {d for vs in adj.values() for d in vs}
        raices = [n for n in todos_nodos if n not in entrantes]
        inicio = raices[0] if raices else todos_nodos[0]

        chain, visited, curr = [], set(), inicio
        while curr and curr not in visited and len(chain) < 10:
            visited.add(curr)
            chain.append(node_label.get(curr, curr[:6]))
            nexts = adj.get(curr, [])
            curr = nexts[0] if nexts else None

        return " → ".join(chain) if chain else "—"

    # Sin relaciones pero con nodos
    if node_label:
        return " → ".join(list(node_label.values())[:10])

    return "—"


def _estado_a_accion(estado: str) -> str:
    return {
        "pendiente": "INICIO_PROCESO",
        "en_proceso": "LLENADO_FORMULARIO",
        "en_progreso": "LLENADO_FORMULARIO",
        "en_revision": "APROBAR",
        "observado": "OBSERVAR",
        "escalado": "ESCALAR",
        "rechazado": "RECHAZAR",
        "finalizado": "FINALIZAR",
        "vencido": "ESCALAR",
    }.get(estado, "LLENADO_FORMULARIO")
